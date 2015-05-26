import time
import constants

class csv:
    def __init__(self, data):
        self.data = data

    def export(self, filename, append=True):
        if filename.split('.')[-1] != '.csv':
            filename += '.csv'
        mode = 'a' if append else 'w'
        with open(filename, mode) as f:
            for row in self.data:
                f.write(','.join(row) + "\n")
        f.close()


class docx:
    def __init__(self, data, fname, text=True):
        self.data = data
        filename = constants.fileloc + fname + \
                   time.asctime().replace(":", "_").replace(" ", "_").replace("%", "_") + ".docx"
        from docx import Document
        document = Document()
        print self.data
        for page in self.data:
            # try:
            title = page[0]
            data = page[1]
            h = document.add_heading(title, 0)
            document.add_paragraph(data)
            document.add_page_break()
            # except:
            #     pass

        document.save(filename)


class text:
    def __init__(self, data):
        self.data = data

    def export(self, filename, append=True):
        if filename.split('.')[-1] != '.txt':
            filename += '.txt'
        mode = 'a' if append else 'w'
        with open(filename, mode) as f:
            for p in self.data:
                f.write(p + "\n")
        f.close()


class json:
    def __init__(self, data):
        self.data = data

    def export(self, filename, append=True):
        import demjson

        js = demjson.encode(self.data)
        if filename.split('.')[-1] != '.json':
            filename += '.json'
        mode = 'a' if append else 'w'
        with open(filename, mode) as f:
            f.write(js)
        f.close()

if __name__ == '__main__':
    import urllib
    from BeautifulSoup import BeautifulSoup as b
    bb = b(urllib.urlopen("http://www.thehindu.com/sci-tech/science/irnss1d-launch/article7043608.ece?homepage=true").read())
    bs = bb.findAll("p")
    docx(bb.find("title").text, "\n".join([s.text for s in bs])).export("a")