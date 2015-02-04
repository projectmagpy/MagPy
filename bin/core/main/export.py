class csv:
	def __init__(self, data):
		self.data = data

	def export(self, filename, append=True):
		if filename.split('.')[-1] != '.csv':
			filename += '.csv'
		mode = 'a' if append else 'w'
		with open(filename, mode) as f:
			for row in self.data:
				f.write(','.join(row) + "\n")
		f.close()


class docx:
	def __init__(self, title, data, text=True):
		self.data = data
		self.title = title
		self.text = text

	def export(self, filename):
		from docx import Document
		if filename.split('.')[-1] != '.docx':
			filename += '.docx'
		document = Document()
		h = document.add_heading(self.title, 0)
		h.add_run('bold').bold = True

		if self.text:
			for p in self.data:
				document.add_paragraph(p)
		else:
			table = document.add_table(rows=len(self.data), cols=len(self.data[0]))
			r, c = -1, -1
			for row in self.data:
				r += 1
				c = 0
				for col in row:
					hdr_cells = table.rows[0].cells
					hdr_cells[c].text = col
					c += 1

		document.add_page_break()
		document.save(filename)


class text:
	def __init__(self, data):
		self.data = data

	def export(self, filename, append=True):
		if filename.split('.')[-1] != '.txt':
			filename += '.txt'
		mode = 'a' if append else 'w'
		with open(filename, mode) as f:
			for p in self.data:
				f.write(p + "\n")
		f.close()

class json:
	def __init__(self, data):
		self.data = data

	def export(self, filename, append=True):
		import demjson
		js = demjson.encode(self.data)
		if filename.split('.')[-1] != '.json':
			filename += '.json'
		mode = 'a' if append else 'w'
		with open(filename, mode) as f:
			f.write(js)
		f.close()